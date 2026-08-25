"""DOCX 输入解析、组装与 CLI 默认导出格式。"""

from __future__ import annotations

import os
import tempfile
import unittest
from unittest.mock import patch

from docx import Document as DocxDocument
from typer.testing import CliRunner

from trans_novel.assemble.docx_writer import _assemble_docx
from trans_novel.assemble.writer import assemble
from trans_novel.cli import _resolve_output_format, app
from trans_novel.config import Config
from trans_novel.ingest.docx_reader import read_docx
from trans_novel.ingest.models import KIND_HEADING, KIND_TEXT
from trans_novel.ingest.segmenter import load_document
from trans_novel.pipeline.runstore import STATUS_DONE, RunStore


def _write_sample_docx(path: str) -> None:
    doc = DocxDocument()
    doc.add_heading("Chapter One", level=1)
    doc.add_paragraph("Hello world.")
    doc.add_heading("Section", level=2)
    doc.add_paragraph("More text.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    doc.save(path)


class TestDocxReader(unittest.TestCase):
    def test_read_headings_paragraphs_and_table(self):
        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "sample.docx")
            _write_sample_docx(path)
            book = read_docx(path, "en", "zh")
        self.assertEqual(book.fmt, "docx")
        self.assertEqual(len(book.chapters), 1)
        self.assertEqual(book.chapters[0].title, "Chapter One")
        kinds = [s.kind for s in book.chapters[0].segments]
        self.assertEqual(kinds[0], KIND_HEADING)
        self.assertEqual(book.chapters[0].segments[0].meta.get("heading_level"), 1)
        self.assertEqual(kinds[1], KIND_TEXT)
        table_segs = [s for s in book.chapters[0].segments if s.meta.get("table_id") == 0]
        self.assertEqual(len(table_segs), 4)
        self.assertEqual(
            {(s.meta["row"], s.meta["col"]): s.source for s in table_segs}[(0, 0)], "A"
        )

    def test_load_document_routes_docx(self):
        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "sample.docx")
            _write_sample_docx(path)
            book = load_document(path, "en", "zh")
        self.assertEqual(book.fmt, "docx")


class TestDocxAssemble(unittest.TestCase):
    def test_assemble_rebuilds_headings_and_table(self):
        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "sample.docx")
            _write_sample_docx(path)
            book = read_docx(path, "en", "zh")
            store = RunStore(os.path.join(directory, "state", "sample"))
            store.save_manifest(
                {
                    "title": "sample",
                    "fmt": "docx",
                    "source_lang": "en",
                    "target_lang": "zh",
                    "source_path": path,
                    "source_sha256": "x",
                    "chapters": [
                        {
                            "index": 0,
                            "title": book.chapters[0].title,
                            "status": STATUS_DONE,
                            "title_translated": "第一章",
                        }
                    ],
                }
            )
            chapter = book.chapters[0]
            for segment in chapter.segments:
                segment.target = f"译:{segment.source}"
            store.save_chapter(chapter)
            out_path = os.path.join(directory, "out.docx")
            written = _assemble_docx(store, out_path)
            self.assertEqual(written, out_path)
            result = DocxDocument(out_path)
            texts = [p.text for p in result.paragraphs if p.text.strip()]
            self.assertIn("译:Chapter One", texts)
            self.assertIn("译:Hello world.", texts)
            self.assertEqual(len(result.tables), 1)
            self.assertEqual(result.tables[0].cell(0, 0).text.strip(), "译:A")
            self.assertEqual(result.tables[0].cell(1, 1).text.strip(), "译:D")

            via_writer = assemble(store, path, out_format="docx")
            self.assertTrue(via_writer.endswith(".docx"))
            self.assertTrue(os.path.isfile(via_writer))


class TestDocxCliDefaults(unittest.TestCase):
    def test_resolve_output_format_defaults(self):
        self.assertEqual(_resolve_output_format("a.docx", None), "docx")
        self.assertEqual(_resolve_output_format("a.epub", None), "epub")
        self.assertEqual(_resolve_output_format("a.docx", "epub"), "epub")
        self.assertEqual(_resolve_output_format("a.txt", "docx"), "docx")

    def test_translate_docx_defaults_out_format(self):
        captured: dict = {}

        class FakeOrchestrator:
            def __init__(self, config, client=None):
                del client
                captured["config"] = config

            def run_all(self, input_path, **kwargs):
                captured["input"] = input_path
                captured["kwargs"] = kwargs
                store = type(
                    "S",
                    (),
                    {"run_dir": "state/sample", "load_usage": staticmethod(lambda: {})},
                )()
                return {
                    "store": store,
                    "outputs": ["output/sample.zh.docx"],
                    "output": "output/sample.zh.docx",
                    "report": {
                        "summary": {
                            "chapters_done": 1,
                            "chapters_total": 1,
                            "terms": 0,
                        }
                    },
                }

        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "sample.docx")
            _write_sample_docx(path)
            with (
                patch(
                    "trans_novel.cli._load_config",
                    return_value=Config.from_dict({"llm": {"provider": "fake"}}),
                ),
                patch("trans_novel.pipeline.orchestrator.Orchestrator", FakeOrchestrator),
            ):
                result = CliRunner().invoke(app, ["translate", path])
        self.assertEqual(result.exit_code, 0, result.output)
        self.assertEqual(captured["kwargs"].get("out_format"), "docx")

    def test_format_docx_is_accepted(self):
        cfg = Config.from_dict({"llm": {"provider": "fake"}})
        captured: dict = {}

        class FakeOrchestrator:
            def __init__(self, config, client=None):
                del client

            def run_all(self, input_path, **kwargs):
                captured["kwargs"] = kwargs
                store = type(
                    "S",
                    (),
                    {"run_dir": "state/sample", "load_usage": staticmethod(lambda: {})},
                )()
                return {
                    "store": store,
                    "outputs": ["output/book.zh.docx"],
                    "output": "output/book.zh.docx",
                    "report": {
                        "summary": {
                            "chapters_done": 1,
                            "chapters_total": 1,
                            "terms": 0,
                        }
                    },
                }

        with tempfile.TemporaryDirectory() as directory:
            path = os.path.join(directory, "book.txt")
            with open(path, "w", encoding="utf-8") as handle:
                handle.write("Hello.\n")
            with (
                patch("trans_novel.cli._load_config", return_value=cfg),
                patch("trans_novel.pipeline.orchestrator.Orchestrator", FakeOrchestrator),
            ):
                result = CliRunner().invoke(app, ["translate", path, "--format", "docx"])
        self.assertEqual(result.exit_code, 0, result.output)
        self.assertEqual(captured["kwargs"].get("out_format"), "docx")


if __name__ == "__main__":
    unittest.main()
