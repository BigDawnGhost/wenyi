"""从 RunStore 重建 Word .docx：标题导航 + 段落样式 + 简易表格。

- ``meta.docx_style``：整段同质，直接套到译文 run（不经 AI）
- ``meta.docx_styles.placements``：混排对齐结果，按 target 偏移切 run
"""

from __future__ import annotations

from typing import Any

from docx import Document as open_docx
from docx.document import Document as DocxDocument
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from ..ingest.models import KIND_HEADING, Chapter
from ..pipeline.runstore import RunStore
from .writer_common import _bilingual_source, _ch_title, _ordered_pair, _seg_text


def _set_outline_level(paragraph, level: int) -> None:
    """确保段落带 outlineLvl，便于 Word 导航窗格。"""
    level = max(1, min(9, level))
    p_pr = paragraph._p.get_or_add_pPr()  # noqa: SLF001
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = p_pr.makeelement(qn("w:outlineLvl"), {})
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level - 1))


def _apply_run_style(run, style: dict[str, Any] | None) -> None:
    """把 meta 中的字符样式应用到 run。"""
    if not style:
        return
    if "bold" in style:
        run.bold = bool(style["bold"])
    if "italic" in style:
        run.italic = bool(style["italic"])
    if style.get("underline"):
        run.underline = True
    size_pt = style.get("size_pt")
    if isinstance(size_pt, (int, float)) and size_pt > 0:
        run.font.size = Pt(float(size_pt))
    color = style.get("color")
    if isinstance(color, str) and len(color) >= 6:
        try:
            run.font.color.rgb = RGBColor.from_string(color[-6:])
        except (ValueError, TypeError):
            pass
    font = style.get("font")
    if isinstance(font, str) and font.strip():
        run.font.name = font.strip()


def _style_slices(
    text: str,
    style: dict[str, Any] | None,
    placements: list[dict[str, Any]] | None,
) -> list[tuple[str, dict[str, Any] | None]]:
    """把文本切成 (fragment, style) 列表；无混排时整段一个切片。"""
    if not text:
        return []
    if not placements:
        return [(text, style)]
    bounds = [0, len(text)]
    usable: list[dict[str, Any]] = []
    for row in placements:
        start = row.get("target_start")
        end = row.get("target_end")
        if not isinstance(start, int) or not isinstance(end, int):
            continue
        start = max(0, min(len(text), start))
        end = max(start, min(len(text), end))
        if start >= end:
            continue
        usable.append({**row, "target_start": start, "target_end": end})
        bounds.extend((start, end))
    if not usable:
        return [(text, style)]
    cuts = sorted(set(bounds))
    slices: list[tuple[str, dict[str, Any] | None]] = []
    for left, right in zip(cuts, cuts[1:]):
        if left >= right:
            continue
        fragment = text[left:right]
        matched: dict[str, Any] | None = None
        for row in usable:
            if row["target_start"] <= left and right <= row["target_end"]:
                matched = {
                    key: row[key]
                    for key in ("bold", "italic", "underline", "color", "size_pt", "font")
                    if key in row
                }
                break
        slices.append((fragment, matched or style))
    return slices or [(text, style)]


def _fill_paragraph(
    paragraph,
    text: str,
    *,
    style: dict[str, Any] | None = None,
    placements: list[dict[str, Any]] | None = None,
    dim: bool = False,
) -> None:
    """清空并按样式切片写入段落。"""
    paragraph.clear()
    for fragment, frag_style in _style_slices(text, style, placements):
        run = paragraph.add_run(fragment)
        _apply_run_style(run, frag_style)
        if dim:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            try:
                run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
            except (AttributeError, ValueError):
                pass


def _add_heading(
    doc: DocxDocument,
    text: str,
    level: int,
    *,
    style: dict[str, Any] | None = None,
    placements: list[dict[str, Any]] | None = None,
) -> None:
    level = max(1, min(9, level))
    style_name = f"Heading {level}"
    try:
        paragraph = doc.add_heading("", level=level)
    except (KeyError, ValueError):
        paragraph = doc.add_paragraph("")
        try:
            paragraph.style = style_name
        except (KeyError, ValueError):
            pass
    _set_outline_level(paragraph, level)
    _fill_paragraph(paragraph, text, style=style, placements=placements)


def _add_normal(
    doc: DocxDocument,
    text: str,
    *,
    style: dict[str, Any] | None = None,
    placements: list[dict[str, Any]] | None = None,
    dim: bool = False,
) -> None:
    paragraph = doc.add_paragraph()
    _fill_paragraph(paragraph, text, style=style, placements=placements, dim=dim)


def _add_bilingual_paragraphs(
    doc: DocxDocument,
    source: str,
    target: str,
    order: str,
    *,
    style: dict[str, Any] | None = None,
    placements: list[dict[str, Any]] | None = None,
) -> None:
    src = _bilingual_source(source, target)
    if not src:
        _add_normal(doc, target, style=style, placements=placements)
        return
    first, second = _ordered_pair(src, target, order)
    if order == "source_first":
        _add_normal(doc, first, dim=False)
        _add_normal(doc, second, style=style, placements=placements, dim=False)
    else:
        _add_normal(doc, first, style=style, placements=placements)
        _add_normal(doc, second, dim=True)


def _segment_style_payload(meta: dict[str, Any]) -> tuple[dict[str, Any] | None, list | None]:
    """返回 (整段样式, 混排 placements)。"""
    uniform = meta.get("docx_style")
    if isinstance(uniform, dict) and uniform:
        return uniform, None
    styles = meta.get("docx_styles")
    if isinstance(styles, dict):
        placements = styles.get("placements")
        if isinstance(placements, list) and placements:
            return None, placements
    return None, None


def _flush_table(
    doc: DocxDocument,
    cells: dict[tuple[int, int], tuple[str, str, dict[str, Any]]],
    rows: int,
    cols: int,
    *,
    bilingual: bool,
    order: str,
) -> None:
    table = doc.add_table(rows=rows, cols=cols)
    try:
        table.style = "Table Grid"
    except (KeyError, ValueError):
        pass
    for r in range(rows):
        for c in range(cols):
            target, source, meta = cells.get((r, c), ("", "", {}))
            style, placements = _segment_style_payload(meta)
            cell = table.cell(r, c)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            if bilingual:
                src = _bilingual_source(source, target)
                if src:
                    first, second = _ordered_pair(src, target, order)
                    _fill_paragraph(paragraph, first, style=style, placements=placements)
                    paragraph.add_run("\n")
                    dim_run = paragraph.add_run(second)
                    dim_run.font.size = Pt(9)
                    dim_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                else:
                    _fill_paragraph(paragraph, target, style=style, placements=placements)
            else:
                _fill_paragraph(paragraph, target, style=style, placements=placements)


def _emit_chapter_blocks(
    doc: DocxDocument,
    chapter: Chapter,
    *,
    bilingual: bool,
    order: str,
) -> None:
    """按段顺序写出；连续同 table_id 聚合成一张表；cont 续段并回上一段。"""
    i = 0
    segs = chapter.segments
    while i < len(segs):
        seg = segs[i]
        meta = seg.meta if isinstance(seg.meta, dict) else {}
        table_id = meta.get("table_id")
        if isinstance(table_id, int):
            cells: dict[tuple[int, int], tuple[str, str, dict[str, Any]]] = {}
            rows = int(meta.get("rows") or 1)
            cols = int(meta.get("cols") or 1)
            while i < len(segs):
                cur = segs[i]
                cur_meta = cur.meta if isinstance(cur.meta, dict) else {}
                if cur_meta.get("table_id") != table_id:
                    break
                r = int(cur_meta.get("row") or 0)
                c = int(cur_meta.get("col") or 0)
                rows = max(rows, int(cur_meta.get("rows") or rows))
                cols = max(cols, int(cur_meta.get("cols") or cols))
                cells[(r, c)] = (_seg_text(cur), cur.source, cur_meta)
                i += 1
            _flush_table(
                doc,
                cells,
                rows,
                cols,
                bilingual=bilingual,
                order=order,
            )
            continue

        if not seg.source.strip() and not (seg.target and seg.target.strip()):
            i += 1
            continue

        target_parts = [_seg_text(seg)]
        source_parts = [seg.source]
        kind = seg.kind
        heading_level = 1
        if kind == KIND_HEADING:
            raw_level = meta.get("heading_level", 1)
            heading_level = raw_level if isinstance(raw_level, int) else 1
        style_meta = meta
        i += 1
        while i < len(segs):
            nxt = segs[i]
            nxt_meta = nxt.meta if isinstance(nxt.meta, dict) else {}
            if not nxt.cont or nxt_meta.get("table_id") is not None:
                break
            target_parts.append(_seg_text(nxt))
            source_parts.append(nxt.source)
            i += 1
        target = "".join(target_parts)
        source = "".join(source_parts)
        style, placements = _segment_style_payload(style_meta)
        if kind == KIND_HEADING:
            _add_heading(doc, target, heading_level, style=style, placements=placements)
        elif bilingual:
            _add_bilingual_paragraphs(
                doc,
                source,
                target,
                order,
                style=style,
                placements=placements,
            )
        else:
            _add_normal(doc, target, style=style, placements=placements)


def _assemble_docx(
    store: RunStore,
    out_path: str,
    *,
    bilingual: bool = False,
    order: str = "target_first",
) -> str:
    """按章节重建 .docx；标题带 outline，样式与表格按 meta 重建。"""
    manifest = store.load_manifest()
    doc = open_docx()
    if doc.paragraphs:
        p0 = doc.paragraphs[0]
        if not p0.text.strip():
            p0.clear()

    first_block = True
    for c in manifest["chapters"]:
        chapter = store.load_chapter(c["index"])
        has_h1 = any(
            s.kind == KIND_HEADING
            and isinstance(s.meta, dict)
            and int(s.meta.get("heading_level") or 1) == 1
            for s in chapter.segments
        )
        title = _ch_title(c)
        if title and not has_h1 and chapter.meta.get("explicit_title"):
            if first_block and doc.paragraphs and not doc.paragraphs[0].text:
                pass
            _add_heading(doc, title, 1)
        _emit_chapter_blocks(doc, chapter, bilingual=bilingual, order=order)
        first_block = False

    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:p") and not (child.text or "").strip():
            texts = [node.text for node in child.iter(qn("w:t")) if node.text]
            if not any(texts) and body.index(child) == 0:
                body.remove(child)
            break

    doc.save(out_path)
    return out_path
